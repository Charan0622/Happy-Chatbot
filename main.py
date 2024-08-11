import os
import streamlit as st
from dotenv import load_dotenv
from fpdf import FPDF
import io
import tempfile
import google.generativeai as gen_ai
from docx import Document
from docx.shared import Pt

# Load environment variables
load_dotenv()

# Configure Streamlit page settings with the dog emoji as the favicon
st.set_page_config(
    page_title="Happy Chat Bot",
    page_icon="🐶",  # Dog emoji favicon
    layout="wide",  # Layout option
    initial_sidebar_state="expanded",  # Sidebar is expanded by default
)

GOOGLE_API_KEY = os.getenv("GOOGLE_API_KEY")

# Set up Google Gemini-Pro AI model
gen_ai.configure(api_key=GOOGLE_API_KEY)
model = gen_ai.GenerativeModel('gemini-pro')

# Function to translate roles between Gemini-Pro and Streamlit terminology
def translate_role_for_streamlit(user_role):
    return "assistant" if user_role == "model" else user_role

# Function to customize and format responses based on personality
def customize_response(response_text, personality):
    if personality == "Friendly 😊":
        return f"😊 <strong>Your Happy ChatBot is Saying in Friendly that:</strong><br><br>{response_text}"
    elif personality == "Professional 👨‍🏫":
        return f"👨‍🏫 <strong>Your Happy ChatBot is Saying in Professionally that:</strong><br><br>{response_text}"
    elif personality == "Humorous 😂":
        return f"😂 <strong>Your Happy ChatBot is Saying in Humorously that:</strong><br><br>{response_text}"
    return response_text

# Custom PDF class with header and footer
class PDF(FPDF):
    def header(self):
        self.set_font('Arial', 'B', 12)
        self.cell(0, 10, 'Chat History', 0, 1, 'C')

    def footer(self):
        self.set_y(-15)
        self.set_font('Arial', 'I', 8)
        self.cell(0, 10, f'Page {self.page_no()}', 0, 0, 'C')

# Function to create a PDF from chat history
def create_pdf(chat_history):
    pdf = PDF()

    # Add fonts
    try:
        pdf.add_font('NotoSansTelugu', '', 'NotoSansTelugu-Regular.ttf', uni=True)
        pdf.add_font('NotoSansTelugu', 'B', 'NotoSansTelugu-Bold.ttf', uni=True)
        pdf.add_font('TiroDevanagari', '', 'TiroDevanagariHindi-Regular.ttf', uni=True)
    except Exception as e:
        print(f"Error loading fonts: {e}")

    pdf.add_page()

    for message in chat_history:
        role = translate_role_for_streamlit(message.role)
        text = f"{role.capitalize()}: {message.parts[0].text}\n"

        # Check if the text contains Telugu or Hindi characters
        if any('\u0C00' <= char <= '\u0C7F' for char in text):
            pdf.set_font('NotoSansTelugu', '', 12)
        elif any('\u0900' <= char <= '\u097F' for char in text):
            pdf.set_font('TiroDevanagari', '', 12)
        else:
            pdf.set_font('Arial', '', 12)
        
        try:
            pdf.multi_cell(0, 10, text)
        except Exception as e:
            print(f"Error writing text to PDF: {e}")

    # Write to a temporary file
    with tempfile.NamedTemporaryFile(delete=False) as tmp_file:
        pdf.output(tmp_file.name)
        tmp_file_path = tmp_file.name

    # Read the content of the temporary file into a BytesIO object
    with open(tmp_file_path, 'rb') as f:
        pdf_output = io.BytesIO(f.read())

    return pdf_output.getvalue()

# Function to create a Word document from chat history
def create_word_doc(chat_history):
    doc = Document()
    doc.add_heading('Chat History', level=1)
    
    for message in chat_history:
        role = translate_role_for_streamlit(message.role)
        text = f"{role.capitalize()}: {message.parts[0].text}"
        
        # Add different styles based on the role
        p = doc.add_paragraph(text)
        p.style.font.size = Pt(12)
    
    # Save to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        doc.save(tmp_file.name)
        tmp_file_path = tmp_file.name

    # Read the content of the temporary file into a BytesIO object
    with open(tmp_file_path, 'rb') as f:
        doc_output = io.BytesIO(f.read())

    return doc_output.getvalue()

# Initialize chat session in Streamlit if not already present
if "chat_session" not in st.session_state:
    st.session_state.chat_session = model.start_chat(history=[])

# Initialize and select personality type
if "personality" not in st.session_state:
    st.session_state.personality = "Friendly"  # Default personality

# Apply custom theme to the page
st.markdown(
    """
    <style>
        body {
            background: linear-gradient(to right, #f5f5f5, #e0e0e0);
            color: #333;
        }
        .stButton>button {
            background-color: #d2b48c;
            color: #fff;
        }
        .stTextInput>input {
            background-color: #f5f5f5;
            color: #333;
        }
        .stMarkdown>p {
            color: #333;
        }
        .chat-message {
            display: flex;
            align-items: flex-start;
            border-radius: 10px;
            padding: 10px;
            margin-bottom: 10px;
            box-shadow: 0 4px 8px rgba(0,0,0,0.1);
        }
        .assistant-message {
            border: 2px solid #d2b48c;
            background-color: #ffffff;
        }
        .user-message {
            border: 2px solid #d2b48c;
            background-color: #f5f5f5;
        }
        .avatar {
            font-size: 1.5em;
            margin-right: 10px;
        }
        .assistant-avatar {
            color: #d2b48c; /* Adjust color for visibility */
        }
        .user-avatar {
            color: #d2b48c; /* Adjust color for visibility */
        }
        .message-content {
            flex: 1;
        }
    </style>
    """, 
    unsafe_allow_html=True
)

# Display the chatbot's title and a brief description
st.title("🐶 Happy Chat Bot")
st.write("🐾 Hello there! 🐾")

st.write("I'm Happy, your friendly dog chatbot! 🐶💬 Though I'm a pup at heart, I'm here to assist you with all your questions. Ask me anything, and I'll do my best to provide loyal and helpful answers! 🐕✨")

st.write("Let's make this chat fun and informative! 🎉.")

# Display personality selection
st.sidebar.header("🐶 ChatBot Personality")
personality_options = ["Friendly 😊", "Professional 👨‍🏫", "Humorous 😂"]
selected_personality = st.sidebar.selectbox("Choose Personality", options=personality_options)
st.session_state.personality = selected_personality

# Display the chat history
for message in st.session_state.chat_session.history:
    role = translate_role_for_streamlit(message.role)
    if role == "assistant":
        st.markdown(
            f"""
            <div class="chat-message assistant-message">
                <div class="avatar assistant-avatar">🐶</div>
                <div class="message-content">{customize_response(message.parts[0].text, st.session_state.personality)}</div>
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            f"""
            <div class="chat-message user-message">
                <div class="avatar user-avatar">👤</div>
                <div class="message-content">{message.parts[0].text}</div>
            </div>
            """,
            unsafe_allow_html=True
        )

# Input field for user's message
user_prompt = st.chat_input("Ask Happy Chat Bot 🐶 ...")
if user_prompt:
    # Add user's message to chat and display it
    st.markdown(
        f"""
        <div class="chat-message user-message">
            <div class="avatar user-avatar">👤</div>
            <div class="message-content">{user_prompt}</div>
        </div>
        """,
        unsafe_allow_html=True
    )

    # Send user's message to Gemini-Pro and get the response
    gemini_response = st.session_state.chat_session.send_message(user_prompt)

    # Customize and display Gemini-Pro's response
    customized_response = customize_response(gemini_response.text, st.session_state.personality)
    st.markdown(
        f"""
        <div class="chat-message assistant-message">
            <div class="avatar assistant-avatar">🐶</div>
            <div class="message-content">{customized_response}</div>
        </div>
        """,
        unsafe_allow_html=True
    )

# Enhance the webpage with additional elements
st.sidebar.header("ChatBot Controls")

# Add dropdown for download options
download_option = st.sidebar.selectbox(
    "Download Chat As",
    ["Select Format", "PDF", "Word Document"]
)

if download_option == "PDF":
    # Create PDF and provide download link
    pdf_data = create_pdf(st.session_state.chat_session.history)
    st.sidebar.download_button(
        label="Download PDF",
        data=pdf_data,
        file_name="chat_history.pdf",
        mime="application/pdf"
    )
elif download_option == "Word Document":
    # Create Word document and provide download link
    word_data = create_word_doc(st.session_state.chat_session.history)
    st.sidebar.download_button(
        label="Download Word Document",
        data=word_data,
        file_name="chat_history.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Display the GIF in the sidebar under the download button
gif_path = "C:/Users/gchar/Scifor_Mini_Project_2/Chatbot.gif"
st.sidebar.image(gif_path, use_column_width=True)  # Adjust the width as needed
